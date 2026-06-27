"""Build a Word (.docx) document from extracted results.

Uses python-docx (already a dependency). Pure-local, no network. Each uploaded
document becomes a section with a filename heading, and each page a "Page N"
heading (annotated with its source / OCR confidence) followed by its text.
"""

import io
import math
import re

# Characters that are illegal in XML 1.0 (and so rejected by python-docx/lxml
# with "All strings must be XML compatible"). We keep the only valid C0 controls
# (tab, newline, carriage return) and strip the rest. A form-feed (U+000C) is the
# common offender: PDF text layers emit it as a page/section separator, so it
# reaches Word export organically and would otherwise 500 the whole export.
_XML_ILLEGAL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _xml_safe(value) -> str:
    """Coerce any value to an XML-safe string for python-docx."""
    if value is None:
        s = ""
    elif isinstance(value, str):
        s = value
    else:
        s = str(value)
    return _XML_ILLEGAL.sub("", s)


def build_docx(documents: list) -> bytes:
    """Render extracted documents to .docx bytes.

    ``documents`` is a list of ``{"filename": str, "pages": [
        {"page": int, "source": "text"|"ocr"|None, "confidence": float|None,
         "text": str}, ...]}``.

    The whole payload is untrusted (it comes from an HTTP body), so this never
    raises on a malformed shape, an XML-illegal control character in any string,
    or a non-finite (NaN/Infinity) confidence — it degrades gracefully instead.
    """
    from docx import Document  # imported lazily so import of this module is cheap

    doc = Document()
    # Be defensive about the posted payload: the body comes from an untrusted
    # HTTP request, so anything that isn't the expected shape (a list of dicts of
    # dicts) is coerced/skipped instead of raising a 500. The frontend always
    # sends well-formed data; a hand-crafted request must not crash the handler.
    if not isinstance(documents, list):
        documents = []
    first = True
    for d in documents:
        if not isinstance(d, dict):
            continue
        if not first:
            doc.add_page_break()
        first = False
        doc.add_heading(_xml_safe(d.get("filename") or "Document"), level=1)

        pages = d.get("pages")
        if not isinstance(pages, list):
            pages = []
        for p in pages:
            if not isinstance(p, dict):
                continue
            source = p.get("source")
            conf = p.get("confidence")
            # Guard finiteness: json.loads accepts bare NaN/Infinity, and
            # round(nan)/round(inf) raise — which would 500 the export.
            if source == "ocr":
                if (
                    isinstance(conf, (int, float))
                    and not isinstance(conf, bool)
                    and math.isfinite(conf)
                ):
                    tag = " — OCR ({}%)".format(round(conf * 100))
                else:
                    tag = " — OCR"
            elif source == "text":
                tag = " — Text layer"
            else:
                tag = ""
            doc.add_heading(_xml_safe("Page {}{}".format(p.get("page"), tag)), level=2)

            text = _xml_safe(p.get("text"))
            if text.strip():
                # Preserve line structure: one paragraph per line.
                for line in text.split("\n"):
                    doc.add_paragraph(line)
            else:
                empty = doc.add_paragraph()
                empty.add_run("(no text on this page)").italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
